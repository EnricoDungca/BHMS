from tkinter import messagebox
from cryptography.fernet import Fernet
from tkinter import filedialog
import mysql.connector
import smtplib
import random as rd
from dotenv import dotenv_values
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import sys, os, datetime, subprocess
import docx

def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller bundle"""
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

# Load .env.secret
env_path = resource_path("Back_End\.env.secret")
if not os.path.exists(env_path):
    messagebox.showerror("Error", f".env.secret file not found at: {env_path}")

config = dotenv_values(env_path)


class Sys_log:
    def __init__(self, Log_name, Log_text):
        self.Log_name = Log_name
        self.Log_text = Log_text
        
    def write_log(self):
        # make Directory
        if not os.path.exists("Logs"):
            os.makedirs("Logs")
        # create log file
        with open(f"Logs/{self.Log_name}.text", "a") as log_file:
            log = {
                "Datetime": datetime.datetime.now().strftime("Date: %Y-%m-%d Time: %H:%M:%S"),
                "LogName": self.Log_name,
                "LogText": self.Log_text
            }
            log_file.write(f"---------------- \n[{log['LogName']}] \n{log['Datetime']} \n{log['LogText']} \n----------------\n\n")
        

class email:
    def __init__(self, email, validation):
        self.email = email
        self.validate = validation
        self.code = rd.randint(100000, 999999)
        
    def verify_gmail(self):
        return self.email.endswith(("@gmail.com", "@icloud.com"))

    def otp_send(self):
        if self.validate and self.verify_gmail():
            try:
                server = smtplib.SMTP("smtp.gmail.com", 587)
                server.starttls()
                server.login(config["EMAIL"], config["GOOGLE_KEY"])
                server.sendmail(config["EMAIL"], self.email, f"""Subject: OTP\n\nYour OTP is {self.code}""")
                server.quit()
                return self.code
            except Exception as e:
                messagebox.showerror("Error", f"Error in sending OTP: {e}. Please Check your internet connection")
        else:
            messagebox.showerror("Error", "Invalid email address")
            
    def send_email(self, message):
        if self.verify_gmail():
            try:
                server = smtplib.SMTP("smtp.gmail.com", 587)
                server.starttls()
                server.login(config["EMAIL"], config["GOOGLE_KEY"])
                server.sendmail(config["EMAIL"], self.email, message.as_string())
                server.quit()
            except Exception as e:
                messagebox.showerror("Error", f"Error in sending email: {e}. Please Check your internet connection")
        else:
            messagebox.showerror("Error", "Invalid email address")
            
class authentication:
    def __init__(self, email, password):
        self.email = email
        self.password = password
        
    def verify_login(self):
        try:
            for row in database_con().read("accounts", "*"):
                if self.email == row[4]:
                    if row[6] == "Active":
                        if self.password == Security().decrypt_str(row[5]):
                            status = row[7] if row[7] == "Active" else "Disabled"
                            return row[0], status, True
            return False, False, False
        except Exception as e:
            messagebox.showerror("Error", str(e))
    
    def verify_login_admin(self):
        try:
            for row in database_con().read("admin", "*"):
                if self.email == row[2] and self.password == row[3]:
                    return row[0], True
            return False
        except Exception as e:
            messagebox.showerror("Error", str(e))
    
    def main(self, mode):
        if mode == "user":
            return self.verify_login()
        elif mode == "admin":
            return self.verify_login_admin()

class Security:
    def __init__(self):
        self.Encrpt_key = config["ENCRYPT_KEY"].encode()
        
    def Encrypt_str(self, text):
        cipher_suite = Fernet(self.Encrpt_key) 
        return cipher_suite.encrypt(text.encode())
    
    def decrypt_str(self, encrypted_str):
        cipher_suite = Fernet(self.Encrpt_key)
        return cipher_suite.decrypt(encrypted_str).decode()

class database_con:
    def __init__(self):
        self.mydb = None
        self.host = config["HOST"]
        self.user = config["USER"]
        self.passw = config["PASSWORD"]
        self.DBName = config["DATABASE"]
    
    def check_connection(self):
        try:
            # If we're in a PyInstaller bundle, make special accommodations
            if getattr(sys, 'frozen', False):
                bundle_dir = sys._MEIPASS
                os.environ['PATH'] = os.path.join(bundle_dir) + os.pathsep + os.environ['PATH']
                
                # Try connecting to XAMPP MySQL without specifying auth_plugin
                try:
                    self.mydb = mysql.connector.connect(
                        host=self.host,
                        user=self.user,
                        password=self.passw,
                        database=self.DBName,
                        use_pure=True  # Force using pure Python implementation
                    )
                    return self.mydb
                except mysql.connector.Error as first_error:
                    # If the first attempt fails, try with explicit auth_plugin
                    try:
                        self.mydb = mysql.connector.connect(
                            host=self.host,
                            user=self.user,
                            password=self.passw,
                            database=self.DBName,
                            auth_plugin='mysql_native_password'
                        )
                        return self.mydb
                    except mysql.connector.Error:
                        # If that also fails, try caching_sha2_password
                        self.mydb = mysql.connector.connect(
                            host=self.host,
                            user=self.user,
                            password=self.passw,
                            database=self.DBName,
                            auth_plugin='caching_sha2_password'
                        )
                        return self.mydb
            else:
                # Standard connection for development environment
                self.mydb = mysql.connector.connect(
                    host=self.host,
                    user=self.user,
                    password=self.passw,
                    database=self.DBName
                )
                return self.mydb
        except mysql.connector.Error as e:
            print("Error: " + str(e))
            messagebox.showerror("Database Connection Error", f"Could not connect to the database: {str(e)}")
            return None
        
    def close_connection(self):
        if self.mydb:
            self.mydb.close()
    
    def insert(self, Table: str, Column_Names: list, values: tuple):
        self.check_connection()
        if not self.mydb:
            messagebox.showwarning("Warning", "No active database connection.")
            return

        cursor = None
        try:
            cursor = self.mydb.cursor()
            if isinstance(Column_Names, str):  
                Column_Names = [Column_Names]

            columns = ", ".join(Column_Names)
            placeholders = ", ".join(["%s"] * len(values))
            sql = f"INSERT INTO {Table} ({columns}) VALUES ({placeholders})"
            cursor.execute(sql, values)
            self.mydb.commit()
            messagebox.showinfo("Success", "Record inserted successfully.")
        except mysql.connector.Error as e:
            messagebox.showerror("Error", f"Error: {e}")
        finally:
            if cursor:
                cursor.close()

    def Record_edit(self, Table, Column_Names, new_value, condition_column, condition_value):
        self.check_connection()
        if not self.mydb:
            messagebox.showwarning("Warning", "No active database connection.")
            return

        cursor = None
        try:
            cursor = self.mydb.cursor()
            sql = f"UPDATE {Table} SET {Column_Names} = %s WHERE {condition_column} = %s"
            cursor.execute(sql, (new_value, condition_value))
            self.mydb.commit()
        except mysql.connector.Error as e:
            messagebox.showerror("Error", "Error: " + str(e))
        finally:
            if cursor:
                cursor.close()

    def Record_delete(self, Tablename, condition_column, condition_value):
        self.check_connection()
        if not self.mydb:
            messagebox.showwarning("Warning", "No active database connection.")
            return

        cursor = None
        try:
            cursor = self.mydb.cursor()
            sql = f"DELETE FROM {Tablename} WHERE {condition_column} = %s"
            cursor.execute(sql, (condition_value,))
            self.mydb.commit()
            messagebox.showinfo("Success", "Record deleted successfully.")
        except mysql.connector.Error as e:
            messagebox.showerror("Error", str(e))
        finally:
            if cursor:
                cursor.close()
    
    def read(self, Tablename, ColumnName):
        self.check_connection()
        if not self.mydb:
            messagebox.showwarning("Warning", "No active database connection.")
            return []

        cursor = None
        try:
            cursor = self.mydb.cursor()
            sql = f"SELECT {ColumnName} FROM {Tablename}"
            cursor.execute(sql)
            result = cursor.fetchall()
            return result if result else []
        except mysql.connector.Error as e:
            messagebox.showerror("Error", "Error: " + str(e))
            return []
        finally:
            if cursor:
                cursor.close()
        
    def column_names(self, Tablename):
        self.check_connection()
        if not self.mydb:
            messagebox.showwarning("Warning", "No active database connection.")
            return []

        cursor = None
        try:
            cursor = self.mydb.cursor()
            cursor.execute(f"SHOW COLUMNS FROM {Tablename}")
            result = cursor.fetchall()
            return result if result else []
        except mysql.connector.Error as e:
            messagebox.showerror("Error", "Error: " + str(e))
            return []
        finally:
            if cursor:
                cursor.close()
    
    def database_backup(self):
        # Validate config keys
        required_keys = ["MYSQLDUMP_PATH", "USER", "DATABASE"]
        for key in required_keys:
            if key not in config:
                print(f"Missing configuration for: {key}")
                return

        # Create backup directory if it doesn't exist
        backup_dir = "DB_Backup"
        os.makedirs(backup_dir, exist_ok=True)

        # Generate timestamped backup filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_filename = f"backup_{timestamp}.sql"
        backup_path = os.path.join(backup_dir, backup_filename)

        # Build mysqldump command
        command = [
            config["MYSQLDUMP_PATH"],
            f"-u{config['USER']}",
            config["DATABASE"]
        ]

        try:
            with open(backup_path, "w", encoding="utf-8") as backup_file:
                subprocess.run(command, stdout=backup_file, stderr=subprocess.PIPE, check=True)
            messagebox.showinfo("Success", f"Database backup created at: {backup_path}")
            
            # Optional: Compress the backup file
            # import gzip, shutil
            # with open(backup_path, 'rb') as f_in, gzip.open(backup_path + '.gz', 'wb') as f_out:
            #     shutil.copyfileobj(f_in, f_out)
            # os.remove(backup_path)
            # print(f"[INFO] Backup compressed to: {backup_path}.gz")

        except subprocess.CalledProcessError as e:
            messagebox.showerror("Error", f"Database backup failed: {e.stderr.decode('utf-8')}")
        except OSError as e:
            messagebox.showerror("Error", f"Database backup failed: {e}")
    
    
class SaveAsDoc:
    def __init__(self, doc_name: str, db_table: str):
        self.doc_name = doc_name
        self.db_table = db_table
        self.timestamp = datetime.datetime.now()
        self.formatted_datetime = self._get_formatted_datetime()
        self.doc = docx.Document()
        
        # Connect to the database once
        db = database_con()
        self.column_names = [col[0] for col in db.column_names(db_table)]
        self.data_rows = db.read(db_table, "*")

    def _get_formatted_datetime(self) -> str:
        return self.timestamp.strftime("Date: %Y-%m-%d    Time: %H:%M:%S")
    
    def write_doc(self):
        try:
            self.doc.add_heading({self.doc_name}, level=0)
            self.doc.add_paragraph(self.formatted_datetime)
            self.doc.add_paragraph('')  # Spacer for clarity

            if not self.data_rows:
                messagebox.showwarning("Warning", "No data found in the database.")
                return
            for row in self.data_rows:
                table = self.doc.add_table(rows=0, cols=2, style='Table Grid')
                for i, col_name in enumerate(self.column_names):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(col_name)
                    row_cells[1].text = str(row[i])
                self.doc.add_paragraph('')  # Spacer between records

            filename = f"{self.doc_name}_{self.timestamp.strftime('%Y%m%d_%H%M%S')}.docx"
            self.doc.save(filename)
            if os.path.exists(filename):
                messagebox.showinfo("Success", f"Document saved as: {filename}")
            else:
                messagebox.showerror("Error", "Failed to save the document.")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")